from io import BytesIO
import datetime
import pandas as pd
from urllib.parse import quote
import datetime
import xlsxwriter
import os
from decimal import Decimal

from django.shortcuts import render
from django.http import FileResponse, HttpResponse, StreamingHttpResponse
from django.utils.encoding import escape_uri_path
from django.conf import settings

from rest_framework import viewsets
from rest_framework import permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status

from .models import *
from .serializers import *
from .filters import *
from goods.pagination import GoodsPagination
from account import permissions as mypermissions
from account.models import AccountModel
from utils import response as myresponse
from goods.models import PriceCycleModel, CategoryModel
from utils.func import is_valid_date
from utils.logger import log_operate

# Create your views here.

class FundsViewset(viewsets.GenericViewSet,
                   myresponse.CustomCreateModelMixin,
                   myresponse.CustomDestroyModelMixin,
                   myresponse.CustomListModelMixin):
    queryset = FundsModel.objects.all()
    serializer_class = FundsModelSerializer

    # 只有教体局组能进行经费来源的添加删除
    def get_permissions(self):
        if self.action in ["create", "destroy"]:
            return [mypermissions.IsRole1()]
        else:
            return [permissions.IsAuthenticated()]


class CartViewset(myresponse.CustomModelViewSet):
    queryset = CartModel.objects.all()
    permission_classes = [mypermissions.IsRole2]

    # 仅能看到由自己账户创建的购物车项
    def get_queryset(self):
        queryset = super().get_queryset()
        user_id = self.request.user.id
        queryset = queryset.filter(creater_id=user_id).order_by('id')
        return queryset
    
    # PATCH方法时使用专用的序列化器
    def get_serializer_class(self):
        if self.action == 'partial_update':
            return CartPatchSerializer
        else:
            return CartModelSerializer
        
    @action(methods=['post'], detail=False)
    def purchase(self, request):
        """
        购物车商品下单
        """
        # 获取购物车项ID列表
        item_list = request.data.get('cart_ids')
        deliver_date = request.data.get('deliver_date')
        note = request.data.get('note')
        user_id = request.user.id

        if not is_valid_date(deliver_date):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        # 判断列表是否为空
        if not item_list:
            return Response({"msg": "未选择购物车商品",
                                "data": None,
                                "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
        
        # 判断送达时间是否为空
        if not deliver_date:
            return Response({"msg": "未选择送达日期",
                                "data": None,
                                "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
        
        # 如果送达日期小于当前日期则返回错误
        deliver_date = datetime.datetime.strptime(deliver_date, '%Y-%m-%d').date()
        if deliver_date < datetime.date.today():
            return Response({"msg": "送达日期不能早于当前日期",
                                "data": None,
                                "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
        
        # 如果当前时间过了中午十二点，则不能下单今天以及明天的订单
        if datetime.datetime.now().hour >= 12:
            if deliver_date == datetime.date.today() or deliver_date == datetime.date.today() + datetime.timedelta(days=1):
                return Response({"msg": "已过下单时间，无法下单今天和明天的订单",
                                "data": None,
                                "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
        
        # 失败列表，存储下单失败的商品的名字
        fail_list = []

        # 过期列表
        expire_list = []

        # 创建成功数，大于0则在失败时不删除订单
        success = 0

        # 商品ID列表
        purchase_list = []

        # 获取送达日期对应的价格周期
        cycle = PriceCycleModel.objects.filter(start_date__lte=deliver_date, end_date__gte=deliver_date).order_by('-start_date').first()
        if not cycle:
            return Response({"msg": "送达日期不在任何价格周期内，无法下单",
                                "data": None,
                                "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)

        # 创建订单实例
        order, is_create = OrdersModel.objects.get_or_create(status=0, creater_id=user_id, deliver_date=deliver_date, cycle=cycle)

        if is_create or order.note is None:
            order.note = note
        else:
            order.note = order.note + ';' + note if note else order.note

        cart_del_list = []
        # 对每一个购物车项，创建一个订单详情实例，并关联在订单实例上
        for item in item_list:
            # 获取购物车项实例和对应商品实例
            try:
                cart = CartModel.objects.get(id=item, creater_id=user_id)
            except:
                if success == 0 and is_create:
                    order.delete()
                return Response({"msg": "购物车项不存在，请刷新后重试",
                                    "data": None,
                                    "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
            try:
                product = cart.product
            except:
                if success == 0 and is_create:
                    order.delete()
                return Response({"msg": "商品不存在，请刷新后重试",
                                    "data": None,
                                    "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)

            if product.status == 0:
                expire_list.append(product.name)
                cart.delete()
                continue
            # 获取商品当前的价格
            price = product.prices.filter(status=2, start_date__lte=deliver_date, end_date__gte=deliver_date).order_by('-id').first()

            # 获取商品用于生成订单详情的图片
            image = product.image
            if image:
                detail_image_path = os.path.join('detail_image', 'goods', image.name.split('/')[-1])
            else:
                detail_image_path = None
            
            # 获取商品用于生成订单详情的资质
            license = product.license
            if license:
                detail_license_path = os.path.join('detail_image', 'license', license.name.split('/')[-1])
            else:
                detail_license_path = None

            # 商品没有可用价格则下单失败，将商品名添加到失败列表中
            # if not price:
            #     fail_list.append(product.name)
            #     continue

            # 创建订单详情实例
            try:
                if not price:
                    OrderDetailModel.objects.create(order=order, product_id=product.id, product_name=product.name, brand=product.brand,
                                            description=product.description, category=product.category.name,
                                            price=0, funds=cart.funds.name, order_quantity=cart.quantity, image=detail_image_path, license=detail_license_path, note=cart.note)
                else:
                    OrderDetailModel.objects.create(order=order, product_id=product.id, product_name=product.name, brand=product.brand,
                                            description=product.description, category=product.category.name,
                                            price=price.price, funds=cart.funds.name, order_quantity=cart.quantity, image=detail_image_path, license=detail_license_path, note=cart.note)
                success += 1
                cart_del_list.append(cart)
                purchase_list.append(product.id)
            except:
                fail_list.append(product.name)
            # 下单后删除购物车项
        

        # 如果全失败，则删除创建的订单实例
        if not order.details:
            order.delete()
            return Response({
                        "msg": "所有商品下单失败",
                        "data": None,
                        "code": status.HTTP_400_BAD_REQUEST}, status=status.HTTP_400_BAD_REQUEST)
        
        # 部分失败则返回失败的商品名
        if fail_list:
            return Response({
                    "msg": "部分商品下单失败",
                    "data": fail_list,
                    "code": status.HTTP_200_OK}, status=status.HTTP_200_OK)
        
        # 如果存在过期商品返回过期商品名
        if expire_list:
            return Response({
                    "msg": "部分商品已下架,刷新以删除商品",
                    "data": expire_list,
                    "code": status.HTTP_200_OK}, status=status.HTTP_200_OK)
        
        # 添加订单的下单总数
        order.product_num = order.details.count()
        order.save()
        for cart in cart_del_list:
            cart.delete()

        # 记录操作
        log_operate(user_id, f"下单{order.id}，商品：{purchase_list}")

        # 返回响应
        return Response({"msg": "商品下单成功",
                    "data": None,
                    "code": status.HTTP_200_OK}, status=status.HTTP_200_OK)
    
class OrdersViewset(viewsets.GenericViewSet,
                   myresponse.CustomListModelMixin,
                   myresponse.CustomDestroyModelMixin,
                   myresponse.CustomUpdateModelMixin):
    queryset = OrdersModel.objects.all()
    serializer_class = OrdersModelSerializer
    filterset_class = OrdersFilter

    def get_queryset(self):
        queryset = super().get_queryset()
        default_order_by = "-deliver_date"
        if self.request.query_params.get("orderby") == "asc":
            default_order_by = "deliver_date"
        # 教体局组和粮油公司组获得所有的订单查看权
        if self.request.user.role == "1" or self.request.user.role == "0":
            return queryset.order_by(default_order_by)
        
        # 学校用户只查询该用户创建的订单
        user_id = self.request.user.id
        queryset = queryset.filter(creater_id=user_id).order_by(default_order_by)
        return queryset

    def get_serializer_class(self):
        if self.action == 'partial_update':
            return OrderPatchSerializer
        else:
            return OrdersModelSerializer
    
    def get_permissions(self):       
        # 仅粮油公司组能进行接单、发货、送达操作
        if self.action in ["accept", "ship", "delivered", "gendeliver", "genfunds", "gendeliverbycat", "destroy", "confirm"]:
            return [mypermissions.IsRole0()]
        elif self.action in ["argue", "agree"]:
            return [mypermissions.IsRole2()]
        elif self.action in ["cancel", "partial_update"]:
            return [mypermissions.IsRole0OrRole2()]
        else:
            return [permissions.IsAuthenticated()]
    
    @action(methods=['get'], detail=True)
    def details(self, request, pk=None):
        """
        获取订单对应的订单详情
        """
        # 获取当前订单实例
        order = self.get_object()

        # 获取订单关联的所有订单详情
        order_details = order.details.all()

        # 制作分页响应
        page_details = self.paginate_queryset(order_details)
        serializer = OrderDetailModelSerializer(page_details, many=True, context={"request": request})
        return self.get_paginated_response(serializer.data)
        # if page_details is not None:
        #     serializer = OrderDetailModelSerializer(page_details, many=True)
        #     return self.get_paginated_response(serializer.data)
        # serializer = OrderDetailModelSerializer(order_details, many=True)
        # return Response({
        #     "msg": "ok",
        #     "data": serializer.data,
        #     "code": status.HTTP_200_OK
        # }, status=status.HTTP_200_OK)
    
    @action(methods=['post'],detail=True)
    def accept(self, request, pk=None):
        """
        修改订单状态为1，表示订单已接单，并添加接单人和接单时间
        """
        order = self.get_object()
                
        # 检查订单当前状态
        if order.status != "0":
            return Response({
                "msg": "接单失败，请检查订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 修改订单数据和状态
        user_id = request.user.id
        order.accepter_id = user_id
        order.accept_time = datetime.datetime.now()
        order.status = 1
        order.save()

        # 记录操作
        log_operate(user_id, f"接单{order.id}")

        # 返回响应
        return Response({
            "msg": "接单成功",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
    
    @action(methods=['post'],detail=True)
    def ship(self, request, pk=None):
        """
        修改订单状态为2，表示订单已发货
        """
        order = self.get_object()

        # 检查订单当前状态
        if order.status != "1":
            return Response({
                "msg": "发货失败，请检查订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
                
        # 修改订单状态
        order.status = 2
        order.save()

        # 记录操作
        log_operate(request.user.id, f"发货{order.id}")

        # 返回响应
        return Response({
            "msg": "发货成功",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
    
    @action(methods=['post'],detail=True)
    def delivered(self, request, pk=None):
        """
        修改订单状态为3，表示订单已送达
        """
        order = self.get_object()

        # 检查订单当前状态
        if order.status != "2":
            return Response({
                "msg": "送达失败，请检查订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 修改订单状态
        order.status = 3
        order.save()

        # 记录操作
        log_operate(request.user.id, f"送达{order.id}")

        # 反回响应
        return Response({
            "msg": "订单送达成功",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
    
    @action(methods=['post'], detail=True)
    def argue(self, request, pk=None):
        """
        对收货情况有异议
        """
        order = self.get_object()

        # 检查订单当前状态
        if order.status != "4":
            return Response({
                "msg": "复核失败，请检查订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 修改订单状态
        order.status = 5
        order.save()

        # 记录操作
        log_operate(request.user.id, f"有异议{order.id}")

        # 放回响应
        return Response({
            "msg": "订单复核结果：订单有疑问",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)

    @action(methods=['post'], detail=True)
    def agree(self, request, pk=None):
        """
        对收货情况没有异议，订单结束
        """
        order = self.get_object()

        # 检查订单当前状态
        if order.status != "4" and order.status != "5":
            return Response({
                "msg": "复核失败，请检查订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 修改订单状态
        order.status = 6
        order.finish_time = datetime.datetime.now()
        order.save()

        # 记录操作
        log_operate(request.user.id, f"订单完成{order.id}")

        # 放回响应
        return Response({
            "msg": "订单复核结果：订单完成",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
    
    @action(methods=['post'], detail=True)
    def confirm(self, request, pk=None):
        """
        对订单确认收货
        """
        order = self.get_object()
        if order.status not in ["3", "4", "5"]:
            return Response({
                "msg": f"当前订单状态为：{order.get_status_display()},无法进行收货操作",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        # err_list存储收货失败的订单详情号
        err_list = []

        # 记录收货ID
        detail_list = []

        # 接收确认收货的订单详情id和收货数量对象
        recv = request.data.get("recv")
        if not recv:
            return Response({
                "msg": "请传入收货情况",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        for data in recv:
            detail_id = data['id']
            received_quantity = data['received_quantity']
            try:
                item = OrderDetailModel.objects.get(id=detail_id)
                # 如果该订单详情不属于该订单，则加入错误列表
                if item.order != order:
                    raise ValueError
                # 当该项之前未有过收货行为时，增加已收货条目数
                if item.received_quantity is None:
                    order.finish_num += 1
                # 更新收货数量和总计价格、时间
                item.received_quantity = received_quantity
                item.cost = float(received_quantity) * float(item.price)
                item.recipient_id = request.user.id
                item.recipient_time = datetime.datetime.now()
                item.save()
                detail_list.append(detail_id)
            except:
                err_list.append(detail_id)


        # 当已收货条目和待收获条目相等时，视为订单全部收货，修改订单状态为待复核
        if order.finish_num == order.product_num:
            order.status = 4
        order.save()

        # 记录操作
        log_operate(request.user.id, f"确认收货{order.id}，订单详情：{detail_list}")

        if len(err_list) != 0:
            return Response({
                "msg": "部分商品收货失败，请检查订单详情号和对应订单号是否正确",
                "data": err_list,
                "code": status.HTTP_200_OK
            }, status=status.HTTP_200_OK)
        
        return Response({
            "msg": "收货成功",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
    
    @action(methods=['post'], detail=False)
    def report(self, request, pk=None):
        """
        根据起止时间和学校ID生成学校在某时间段内的订单报表
        """

        # 获取data中的参数
        start_date = request.data.get("start_date")
        end_date = request.data.get("end_date")
        school_id = request.data.get("school_id", None)
        queryset = self.get_queryset()

        # 判断起止时间是否合法
        if not is_valid_date(start_date) or not is_valid_date(end_date):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        
        # 没传入shool_id时表示是学校账户访问自己的报表
        if school_id is None:
            # 判断如果账户是非学校账户，则返回当期所有订单详情
            if request.user.role in ["0", "1"]:
                # 过滤在起止日期内的订单详情
                # _end_date = datetime.datetime.strptime(end_date, "%Y-%m-%d") + datetime.timedelta(days=1)
                queryset = queryset.filter(deliver_date__gte=start_date, deliver_date__lte=end_date).order_by('creater_id', 'id')

            # 如果是学校账户，则返回当期用户的订单详情
            else:
                # 获取当天的最后时间
                # _end_date = datetime.datetime.strptime(end_date, "%Y-%m-%d") + datetime.timedelta(days=1)

                # 过滤在起止日期内的订单详情
                queryset = queryset.filter(deliver_date__gte=start_date, deliver_date__lte=end_date).order_by('id')
        
        # 如果传入的school_id不是学校账户则返回错误
        elif AccountModel.objects.filter(id=school_id).first().role != '2':
            return Response({
                "msg": "访问对象非学校",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 传入了school_id，表示教体局或粮油公司访问报表
        else:
            # 判断当期是否是教体局或粮油公司组用户
            if request.user.role not in ["0", "1"]:
                return Response({
                    "msg": "没有访问权限",
                    "data": None,
                    "code": status.HTTP_401_UNAUTHORIZED
                }, status=status.HTTP_401_UNAUTHORIZED)
            
            # 过滤某一学校在起止日期内的订单详情
            # _end_date = datetime.datetime.strptime(end_date, "%Y-%m-%d") + datetime.timedelta(days=1)
            queryset = queryset.filter(deliver_date__gte=start_date, deliver_date__lte=end_date, creater_id=school_id).order_by('id')


        # 如果查询结果为空，表示时间段内没有订单
        if not queryset.exists():
            return Response({
                "msg": "所选时间段没有订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 制作表格数据
        data = []
        for order in queryset:
            school_name = AccountModel.objects.get(id=order.creater_id).first_name
            order_details = order.details.all()
            for detail in order_details:
                detail_data = {
                    "商品编号": detail.product_id,
                    "订单编号": order.id,
                    "商品名称": detail.product_name,
                    "商品规格": detail.description,
                    "商品品牌": detail.brand,
                    "商品种类": detail.category,
                    "商品单价": detail.price,
                    "经费来源": detail.funds,
                    "订购数量": detail.order_quantity,
                    "实收数量": detail.received_quantity,
                    "总价": detail.cost,
                    "下单时间": str(order.create_time).split(".")[0],
                    "送货日期": str(order.deliver_date),
                    "收货时间": str(detail.recipient_time).split(".")[0] if detail.recipient_time else None,
                    "学校": school_name
                }
                data.append(detail_data)

        # 返回文件响应
        df = pd.DataFrame(data)
        file_name = f"{start_date}~{end_date}_订单报表.xlsx"

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{quote(file_name)}"'

        with pd.ExcelWriter(response, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')

        return response

    @action(methods=['post'], detail=False)
    def gendeliver(self, request, pk=None):
        """
        生成对应日期的送货表
        """
        deliver_date = request.data.get("deliver_date")
        school_id = request.data.get("school_id")


        if not deliver_date:
            return Response({
                "msg": "未选择送货时间",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        

        if not is_valid_date(deliver_date):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not school_id:
            return Response({
                "msg": "未选择收货单位",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取指定送货时间和收货单位的查询集
        queryset = self.get_queryset()
        queryset = queryset.filter(deliver_date=deliver_date, creater_id=school_id)
        if not queryset:
            return Response({
                "msg": "未找到对应订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 制作表单数据
        category_data = {}
        count = {}
        categorys = ["粮油类", "其他类"]

        # 为每个种类制作data列表
        for category in categorys:
            category_data[category] = []
            count[category] = 1

        # 制作备注
        note_list = []

        for order in queryset:
            # 跳过已完成和未接单的订单
            if order.status == '6' or order.status == '0':
                continue
            if order.note:
                note_list.append(f"{order.note}")
            details = order.details.all()
            for detail in details:
                category = "其他类" if detail.category != "粮油类" else "粮油类"

                detail_data = []
                detail_data.append(count[category])
                count[category] += 1
                detail_data.append(detail.product_name)
                detail_data.append(detail.brand)
                detail_data.append(detail.description)
                detail_data.append(detail.order_quantity)
                detail_data.append('')
                detail_data.append(detail.note)
                category_data[category].append(detail_data)

        note = ";".join(note_list)
        # data = []
        # no = 1
        # for order in queryset:
        #     if order.status == '6':
        #         continue
        #     details = order.details.all()
        #     for detail in details:
        #         detail_data = []
        #         detail_data.append(no)
        #         no += 1
        #         detail_data.append(detail.product_name)
        #         detail_data.append(detail.brand)
        #         detail_data.append(detail.description)
        #         detail_data.append(detail.order_quantity)
        #         detail_data.append('')
        #         detail_data.append('')
        #         data.append(detail_data)
        
        # 填充表格数据
        # data1 = [
        #     [1, '产品A', '品牌X', '规格1', '个', 100, 10, 1000, ''],
        #     [2, '产品B', '品牌Y', '规格2', '盒', 50, 20, 1000, ''],
        #     # 添加更多数据...
        # ]

        # 获取收货单位
        first_name = AccountModel.objects.get(id=school_id).first_name
        
        """
        以下为生成表格样式
        """
        output = BytesIO()

        workbook = xlsxwriter.Workbook(output, {"in_memory":True})

        # 表示种类为空的数量，如果所有种类的数据都为空，表示所选时间没有要送的订单
        empty_num = 0

        # 遍历字典，为每个种类创建一个data列表，存储该种类下的订单详情信息，用于生成表格
        for k, data in category_data.items():
            if len(data) == 0:
                empty_num += 1
                continue
            worksheet = workbook.add_worksheet(name=k)

            # 设置列宽度
            worksheet.set_column('A:G', 9)   # 单位列

            # 合并单元格并添加标题
            worksheet.set_row(0, 33)
            worksheet.merge_range('A1:G1', '泸定县粮油购销有限责任公司配送清单', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 18, 
                'bold': True
            }))

            worksheet.set_row(1, 25)
            worksheet.write('A2', '收货单位', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range('B2:D2', first_name, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.write('E2', '配送日期', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range('F2:G2', deliver_date, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))

            # 添加表头
            worksheet.set_row(2, 25)
            headers = ['行号', '品名', '品牌', '规格', '预订数量', '实收数量', '备注']
            worksheet.write_row('A3', headers, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
                'border': 1
            }))

            # 写入数据
            row_start = 3
            row_for_total = 0
            for row, record in enumerate(data, start=row_start):
                worksheet.set_row(row, 25)
                worksheet.write_row(row, 0, record, workbook.add_format({
                    'border': 1,
                    'font_size': 11,
                    'align': 'center',
                    'valign': 'vcenter'
                    }))
                row_for_total = row

            # 添加备注：
            row_for_total += 1
            worksheet.set_row(row_for_total, 25)
            worksheet.write(f'A{row_for_total+1}', '备注', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range(f'B{row_for_total+1}:G{row_for_total+1}', f'{note}', workbook.add_format({
                'align': 'left', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            
            # 添加额外信息
            row_for_total += 1
            worksheet.set_row(row_for_total, 25)
            worksheet.write(f'A{row_for_total+1}', '送货人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range(f'B{row_for_total+1}:C{row_for_total+1}', '')

            worksheet.write(f'D{row_for_total+1}', '验收人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            # worksheet.merge_range(f'E{row_for_total+1}:F{row_for_total+1}', '')

            worksheet.write(f'F{row_for_total+1}', '负责人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            # worksheet.merge_range(f'H{row_for_total+1}:I{row_for_total+1}', '')

        # 关闭Excel文件
        workbook.close()

        # 判单是否数据全为空
        if len(category_data) == empty_num:
            return Response({
                "msg":"所选日期没有待送货物品",
                "data":None,
                "code":status.HTTP_204_NO_CONTENT
            }, status=status.HTTP_204_NO_CONTENT)

        output.seek(0)
        response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        file_name = f"{first_name}_{deliver_date}_送货单.xlsx"
        response['Content-Disposition'] = f'attachment; filename="{quote(file_name)}"'

        return response

    @action(methods=['post'], detail=False)
    def genfunds(self, request, pk=None):
        from docx import Document
        from docx.shared import Pt  # 用于设置字体大小
        from docx.oxml.ns import qn  # 用于设置字体

        # 定位模板文件
        output = BytesIO()
        template_path = os.path.join(settings.MEDIA_ROOT, 'funds_template.docx')
        doc = Document(template_path)

        # 接收参数
        school_id = request.data.get("school_id")
        start_date = request.data.get("start_date")
        end_date = request.data.get("end_date")

        if not is_valid_date(start_date) or not is_valid_date(end_date):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        if not school_id:
            return Response({
                "msg": "请传入学校ID",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        if not start_date or not end_date:
            return Response({
                "msg": "请传入起止时间",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 生成填充数据
        try:
            school = AccountModel.objects.get(id=school_id)
        except:
            return Response({
                "msg": "ID对应的学校不存在",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if school.role != "2":
            return Response({
                "msg": "ID对应的学校不存在",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取学校名称
        school_name = school.first_name
        

        # 获取查询集
        queryset = self.get_queryset()
        queryset = queryset.filter(deliver_date__gte=start_date, deliver_date__lte=end_date, creater_id=school_id, status=6)

        if not queryset:
            return Response({
                "msg": "未查询到时间段内订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        # 生成日期表示
        start_date = datetime.datetime.strptime(start_date, "%Y-%m-%d")
        end_date = datetime.datetime.strptime(end_date, "%Y-%m-%d")
        start_date2 = start_date.strftime("%Y 年 %m 月 %d 日")
        end_date2 = end_date.strftime("%Y 年 %m 月 %d 日")
        start_date = start_date.strftime("%Y-%m-%d")
        end_date = end_date.strftime("%Y-%m-%d")
        date = f"{start_date2}  ------------ {end_date2}"

        # 获取经费情况
        funds_dic = {}
        for order in queryset:
            order_details = order.details.all()
            for detail in order_details:
                now_fund = detail.funds
                funds_dic[now_fund] = funds_dic.get(now_fund, 0) + detail.cost
        
        total = 0
        for k,v in funds_dic.items():
            total += v

        """
        以下为生成docx
        """
        # 遍历文档中的所有表格
        for table in doc.tables:
            for row in table.rows:
                if row.cells[0].text == "学校名称":
                    row.cells[1].text = ''
                    run = row.cells[1].paragraphs[0].add_run(school_name)
                    run.font.name = '宋体'  # 设置字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 强制设置中文字体为宋体
                    run.font.size = Pt(11)  # 设置字体大小为11号
                if row.cells[0].text == "结算时间":
                    row.cells[1].text = ''
                    run = row.cells[1].paragraphs[0].add_run(date)
                    run.font.name = '宋体'  # 设置字体
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 强制设置中文字体为宋体
                    run.font.size = Pt(11)  # 设置字体大小为11号
                if row.cells[0].text == "结算资金（元）":
                    if row.cells[1].text == "以上四项合计金额：":
                        res = row.cells[1].text + str(total)
                        row.cells[1].text = ''
                        run = row.cells[1].paragraphs[0].add_run(res)
                        run.font.name = '宋体'  # 设置字体
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 强制设置中文字体为宋体
                        run.font.size = Pt(11)  # 设置字体大小为11号
                    else:
                        # 使用经费字段名和模板进行匹配
                        for k, v in funds_dic.items():
                            if k in row.cells[1].text:
                                res = row.cells[1].text + str(v)
                                row.cells[1].text = ''
                                run = row.cells[1].paragraphs[0].add_run(res)
                                run.font.name = '宋体'  # 设置字体
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 强制设置中文字体为宋体
                                run.font.size = Pt(11)  # 设置字体大小为11号
        
        doc.save(output)
        output.seek(0)
        response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        file_name = f"{start_date}~{end_date}_{school_name}经费情况.docx"
        response['Content-Disposition'] = f'attachment; filename={quote(file_name)}'


        return response
    
    @action(methods=['post'], detail=True)
    def cancel(self, request, pk=None):
        """
        用于取消订单。当学校用户取消订单时，仅能取消未接单的订单；当粮油公司取消订单时，仅能取消待发货的订单
        """
        order = self.get_object()

        # 学校用户取消订单
        if self.request.user.role == "2":
            if order.status != "0":
                return Response({
                    "msg": "订单已接单，无法取消",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)

            # 取消订单后将订单详情中的商品重新加入购物车
            for detail in order.details.all():
                product = GoodsModel.objects.get(id=detail.product_id)
                funds = FundsModel.objects.get(name=detail.funds)
                        
                # 检查购物车中是否有相同人创建的相同的商品和经费来源的项，如果有，进行累加
                existing_cart = CartModel.objects.filter(product=product, funds=funds, creater_id=order.creater_id).first()
                if existing_cart:
                    existing_cart.quantity += int(detail.order_quantity)
                    existing_cart.save()
                else:
                    # 没有则创建新的购物车实例
                    CartModel.objects.create(product=product, funds=funds, quantity=detail.order_quantity, creater_id=order.creater_id, note=detail.note)

            # 记录操作
            log_operate(request.user.id, f"取消订单{order.id}")
            order.delete() 

            return Response({
                "msg": "订单取消成功",
                "data": None,
                "code": status.HTTP_200_OK
            }, status=status.HTTP_200_OK)


        elif self.request.user.role == "0":
            if order.status != "1":
                return Response({
                    "msg": "订单已发货，无法取消",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # 将订单的状态设置为0，表示订单未接单
            order.status = 0
            order.accepter_id = None
            order.accept_time = None
            order.save()

            # 记录操作
            log_operate(request.user.id, f"取消订单{order.id}")

            return Response({
                "msg": "订单取消接单成功",
                "data": None,
                "code": status.HTTP_200_OK
            }, status=status.HTTP_200_OK)
        
    @action(methods=['post'], detail=False)
    def gendeliverbycat(self, request, pk=None):
        """
        生成按照种类分类的送货表
        """
        deliver_date = request.data.get("deliver_date")
        category_list = request.data.get("category_list")
        school_id = request.data.get("school_id")

        if not deliver_date:
            return Response({
                "msg": "未选择送货时间",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
    

        if not is_valid_date(deliver_date):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)

        
        if not school_id:
            return Response({
                "msg": "未选择收货单位",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not category_list:
            return Response({
                "msg": "未选择商品种类",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取指定送货时间和收货单位的查询集
        queryset = self.get_queryset()
        queryset = queryset.filter(deliver_date=deliver_date, creater_id=school_id)
        if not queryset:
            return Response({
                "msg": "未找到对应订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        
        # 制作表单数据
        category_data = {}
        count = {}
        categorys = []

        for category_id in category_list:
            category = CategoryModel.objects.filter(id=category_id).first()
            if not category:
                return Response({
                    "msg": f"id为{category_id}的商品种类不存在",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)
            categorys.append(category.name)

        # 为每个种类制作data列表
        for category in categorys:
            category_data[category] = []
            count[category] = 1

        # 制作备注
        note_list = []

        for order in queryset:
            # 跳过已完成和未接单的订单
            if order.status == '6' or order.status == '0':
                continue
            if order.note:
                note_list.append(f"{order.note}")
            details = order.details.all()
            for detail in details:
                category = detail.category
                if category not in categorys:
                    continue

                detail_data = []
                detail_data.append(count[category])
                count[category] += 1
                detail_data.append(detail.product_name)
                detail_data.append(detail.brand)
                detail_data.append(detail.description)
                detail_data.append(detail.order_quantity)
                detail_data.append('')
                detail_data.append(detail.note)
                category_data[category].append(detail_data)

        note = ";".join(note_list)
        

        # 获取收货单位
        first_name = AccountModel.objects.get(id=school_id).first_name
        
        """
        以下为生成表格样式
        """
        output = BytesIO()

        workbook = xlsxwriter.Workbook(output, {"in_memory":True})

        # 表示种类为空的数量，如果所有种类的数据都为空，表示所选时间没有要送的订单
        empty_num = 0

        # 遍历字典，为每个种类创建一个data列表，存储该种类下的订单详情信息，用于生成表格
        for k, data in category_data.items():
            if len(data) == 0:
                empty_num += 1
                continue
            worksheet = workbook.add_worksheet(name=k)

            # 设置列宽度
            worksheet.set_column('A:G', 9)   # 单位列

            # 合并单元格并添加标题
            worksheet.set_row(0, 33)
            worksheet.merge_range('A1:G1', '泸定县粮油购销有限责任公司配送清单', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 18, 
                'bold': True
            }))

            worksheet.set_row(1, 25)
            worksheet.write('A2', '收货单位', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range('B2:D2', first_name, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.write('E2', '配送日期', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range('F2:G2', deliver_date, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))

            # 添加表头
            worksheet.set_row(2, 25)
            headers = ['行号', '品名', '品牌', '规格', '预订数量', '实收数量', '备注']
            worksheet.write_row('A3', headers, workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
                'border': 1
            }))

            # 写入数据
            row_start = 3
            row_for_total = 0
            for row, record in enumerate(data, start=row_start):
                worksheet.set_row(row, 25)
                worksheet.write_row(row, 0, record, workbook.add_format({
                    'border': 1,
                    'font_size': 11,
                    'align': 'center',
                    'valign': 'vcenter'
                    }))
                row_for_total = row

            # 添加备注：
            row_for_total += 1
            worksheet.set_row(row_for_total, 25)
            worksheet.write(f'A{row_for_total+1}', '备注', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range(f'B{row_for_total+1}:G{row_for_total+1}', f'{note}', workbook.add_format({
                'align': 'left', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            
            # 添加额外信息
            row_for_total += 1
            worksheet.set_row(row_for_total, 25)
            worksheet.write(f'A{row_for_total+1}', '送货人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            worksheet.merge_range(f'B{row_for_total+1}:C{row_for_total+1}', '')

            worksheet.write(f'D{row_for_total+1}', '验收人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            # worksheet.merge_range(f'E{row_for_total+1}:F{row_for_total+1}', '')

            worksheet.write(f'F{row_for_total+1}', '负责人', workbook.add_format({
                'align': 'center', 
                'valign': 'vcenter', 
                'font_size': 11, 
            }))
            # worksheet.merge_range(f'H{row_for_total+1}:I{row_for_total+1}', '')

        # 关闭Excel文件
        workbook.close()

        # 判单是否数据全为空
        if len(category_data) == empty_num:
            return Response({
                "msg":"所选日期没有待送货物品",
                "data":None,
                "code":status.HTTP_204_NO_CONTENT
            }, status=status.HTTP_204_NO_CONTENT)

        output.seek(0)
        response = HttpResponse(output, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        file_name = f"{first_name}_{deliver_date}_送货单.xlsx"
        response['Content-Disposition'] = f'attachment; filename="{quote(file_name)}"'

        return response

    @action(methods=['post'], detail=True)
    def addproduct(self, request, pk =None):
        order = self.get_object()

        product_id = request.data.get("product_id")
        funds_id = request.data.get("funds_id")
        quantity = request.data.get("quantity")
        note = request.data.get("note")

        if self.request.user.role == "2":
            if order.status != "0":
                return Response({
                    "msg": "订单已接单，无法添加商品",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)
            
        if self.request.user.role == "0":
            if order.status == "0":
                return Response({
                    "msg": "无法对未接单的订单添加商品",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)

        if not product_id or not funds_id or not quantity:
            return Response({
                "msg": "商品ID、经费来源ID、数量不能为空",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取商品实例
        product = GoodsModel.objects.filter(id=product_id).first()
        if not product:
            return Response({
                "msg": "商品不存在",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取商品的经费来源实例
        funds = FundsModel.objects.filter(id=funds_id).first()
        if not funds:
            return Response({
                "msg": "经费来源不存在",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取商品的价格
        price = product.prices.filter(status=2, start_date__lte=order.deliver_date, end_date__gte=order.deliver_date).order_by('-id').first()

        # 获取商品用于生成订单详情的图片
        image = product.image
        if image:
            detail_image_path = os.path.join('detail_image', 'goods', image.name.split('/')[-1])
        else:
            detail_image_path = None
        
        # 获取商品用于生成订单详情的资质
        license = product.license
        if license:
            detail_license_path = os.path.join('detail_image', 'license', license.name.split('/')[-1])
        else:
            detail_license_path = None

        # 创建订单详情
        try:
            if not price:
                OrderDetailModel.objects.create(order=order, product_id=product.id, product_name=product.name, brand=product.brand,
                                        description=product.description, category=product.category.name,
                                        price=0, funds=funds.name, order_quantity=quantity, image=detail_image_path, license=detail_license_path, note=note)
            else:
                OrderDetailModel.objects.create(order=order, product_id=product.id, product_name=product.name, brand=product.brand,
                                        description=product.description, category=product.category.name,
                                        price=price.price, funds=funds.name, order_quantity=quantity, image=detail_image_path, license=detail_license_path, note=note)
        except:
            return Response({
                "msg": "添加商品失败",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 订单商品数量加1
        order.product_num += 1
        order.save()

        # 记录操作
        log_operate(request.user.id, f"添加商品{product.id}到订单{order.id}")

        return Response({
            "msg": "添加商品成功",
            "data": None,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)

    def perform_destroy(self, instance):
        if instance.status == "0":
            return Response({
                "msg": "无法删除未接单的订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 记录操作
        log_operate(self.request.user.id, f"删除订单{instance.id}")

        return super().perform_destroy(instance)

    def perform_update(self, serializer):
        if not is_valid_date(str(serializer.validated_data.get("deliver_date"))):
            return Response({
                "msg": "日期格式错误",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if self.request.user.role == "2" and serializer.instance.status not in ["0", "1"]:
            return Response({
                "msg": "无法修改已发货的订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if self.request.user.role == "0" and serializer.instance.status == "0":
            return Response({
                "msg": "无法修改未接单的订单",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 记录操作
        log_operate(self.request.user.id, f"修改订单{serializer.instance.id}送货日期为{serializer.validated_data.get('deliver_date')}")
        return super().perform_update(serializer)

    @action(methods=['post'], detail=False)
    def gentotal(self, request, pk=None):
        school_id = request.data.get("school_id")
        funds_id = request.data.get("funds_id")
        cycle_id = request.data.get("cycle_id")
        status_list = request.data.get("status_list")
        
        if not funds_id:
            return Response({
                "msg": "未选择经费来源",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not cycle_id:
            return Response({
                "msg": "未选择周期",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        if not status_list:
            return Response({
                "msg": "未选择订单状态",
                "data": None,
                "code": status.HTTP_400_BAD_REQUEST
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # 获取查询集
        queryset = self.get_queryset()

        if self.request.user.role != "2" and school_id:
            queryset = queryset.filter(creater_id=school_id)

        queryset = queryset.filter(cycle=cycle_id, status__in=status_list)

        if not queryset:
            return Response({
                "msg": "未找到对应订单",
                "data": None,
                "code": status.HTTP_404_NOT_FOUND
            }, status=status.HTTP_404_NOT_FOUND)
        
        total = 0

        for order in queryset:
            for detail in order.details.all():
                if detail.funds == FundsModel.objects.filter(id=funds_id).first().name:
                    # total += detail.cost if detail.cost else Decimal(detail.price) * Decimal(detail.order_quantity)
                    if detail.cost is not None:
                        total += detail.cost
                    else:
                        total += Decimal(detail.price) * Decimal(detail.order_quantity)
        
        return Response({
            "msg": "获取成功",
            "data": total,
            "code": status.HTTP_200_OK
        }, status=status.HTTP_200_OK)
        




class OrderDetailsViewset(viewsets.GenericViewSet,
                          myresponse.CustomDestroyModelMixin):
    """
    仅用于确认收货，无CRUD方法
    """
    queryset = OrderDetailModel.objects.all()
    permission_classes = [mypermissions.IsRole0OrRole2]
    serializer_class = OrderDetailModelSerializer

    # 仅获取当期用户创建的订单详情
    def get_queryset(self):
        queryset = super().get_queryset()
        if self.request.user.role == "2":
            queryset = queryset.filter(order__creater_id=self.request.user.id)
        return queryset
    
    def perform_destroy(self, instance):
        order_obj = instance.order
        # 如果当前为学校用户，只能删除未接单的订单详情
        if self.request.user.role == "2":
            if order_obj.status != "0":
                return Response({
                    "msg": "订单已接单，无法删除",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)
            
        if self.request.user.role == "0":
            if order_obj.status == "0":
                return Response({
                    "msg": "无法对未接单的订单删除商品",
                    "data": None,
                    "code": status.HTTP_400_BAD_REQUEST
                }, status=status.HTTP_400_BAD_REQUEST)
            
        # 删除时，将order的product_num减少1
        order_obj.product_num -= 1
        if instance.received_quantity:
            order_obj.finish_num -= 1
        order_obj.save()

        # 记录操作
        log_operate(self.request.user.id, f"删除订单号{order_obj.id}的订单详情{instance.id}")

        # 如果订单详情的商品数量为0，则删除订单
        if order_obj.product_num == 0:
            order_obj.delete()
            # 记录操作
            log_operate(self.request.user.id, f"删除订单{order_obj.id}")
        
        return super().perform_destroy(instance)
    
    # @action(methods=['post'], detail=True)
    # def confirm(self, request, pk=None):
    #     """
    #     确认收货，并修改订单的完成项和完成时间
    #     """
    #     # 获取data中的数据，实际收货数量
    #     received_quantity = request.data.get("received_quantity")

    #     # 获取到当期订单详情实例，用于修改
    #     item = self.get_object()

    #     # 修改订单详情数据
    #     item.received_quantity = received_quantity
    #     item.cost = float(received_quantity) * float(item.price)
    #     item.recipient_id = request.user.id
    #     item.recipient_time = datetime.datetime.now()
    #     item.save()

    #     # 修改订单详情关联的订单的数据，修改完成项
    #     order = item.order
    #     order.finish_num += 1

    #     # 如果完成项等于下单数量，则添加完成时间，表示订单已全部完成
    #     if order.finish_num == order.product_num:
    #         order.status = 4
    #     order.save()

    #     # 返回响应
    #     return Response({
    #         "msg": "商品收货成功",
    #         "data": None,
    #         "code": status.HTTP_200_OK
    #     }, status=status.HTTP_200_OK)
    


    
